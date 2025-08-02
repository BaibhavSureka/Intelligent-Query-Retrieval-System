import asyncio
import aiofiles
import requests
import logging
import uuid
import re
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO
from email.parser import Parser
from email_reply_parser import EmailReplyParser
from urllib.parse import urlparse
import hashlib

# Document processing libraries
import PyPDF2
import fitz  # PyMuPDF for better PDF processing
from docx import Document
from bs4 import BeautifulSoup

from models import DocumentChunk
from config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of various document formats"""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.max_size = settings.max_document_size
        
    async def process_document_from_url(self, document_url: str) -> List[DocumentChunk]:
        """
        Download and process a document from URL
        
        Args:
            document_url: URL to the document
            
        Returns:
            List of DocumentChunk objects
        """
        try:
            # Download document
            logger.info(f"Downloading document from: {document_url}")
            content = await self._download_document(document_url)
            
            # Determine document type
            doc_type = self._detect_document_type(document_url, content)
            logger.info(f"Detected document type: {doc_type}")
            
            # Process based on type
            if doc_type == 'pdf':
                text_content = await self._process_pdf(content)
            elif doc_type == 'docx':
                text_content = await self._process_docx(content)
            elif doc_type == 'email':
                text_content = await self._process_email(content)
            else:
                # Try to process as plain text
                text_content = content.decode('utf-8', errors='ignore')
            
            # Create chunks
            chunks = self._create_chunks(text_content, document_url)
            logger.info(f"Created {len(chunks)} chunks from document")
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document from {document_url}: {str(e)}")
            raise
    
    async def _download_document(self, url: str) -> bytes:
        """Download document from URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            if len(response.content) > self.max_size:
                raise ValueError(f"Document size exceeds maximum allowed size of {self.max_size} bytes")
            
            return response.content
            
        except requests.RequestException as e:
            logger.error(f"Failed to download document from {url}: {str(e)}")
            raise
    
    def _detect_document_type(self, url: str, content: bytes) -> str:
        """Detect document type from URL and content"""
        # Check URL extension
        parsed_url = urlparse(url)
        path = parsed_url.path.lower()
        
        if path.endswith('.pdf'):
            return 'pdf'
        elif path.endswith('.docx') or path.endswith('.doc'):
            return 'docx'
        elif path.endswith('.eml') or path.endswith('.msg'):
            return 'email'
        
        # Check content magic bytes
        if content.startswith(b'%PDF'):
            return 'pdf'
        elif content.startswith(b'PK\x03\x04') and b'word/' in content:
            return 'docx'
        elif b'From:' in content[:1000] or b'To:' in content[:1000]:
            return 'email'
        
        return 'text'
    
    async def _process_pdf(self, content: bytes) -> str:
        """Extract text from PDF using both PyPDF2 and PyMuPDF for better coverage"""
        text_content = ""
        
        try:
            # First try with PyMuPDF (more reliable)
            pdf_document = fitz.open(stream=content, filetype="pdf")
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    text_content += f"\n--- Page {page_num + 1} ---\n{text}\n"
            
            pdf_document.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content += f"\n--- Page {page_num + 1} ---\n{text}\n"
                        
            except Exception as e2:
                logger.error(f"Both PDF processors failed: {str(e2)}")
                raise
        
        if not text_content.strip():
            raise ValueError("No text content could be extracted from PDF")
        
        return self._clean_text(text_content)
    
    async def _process_docx(self, content: bytes) -> str:
        """Extract text from DOCX document"""
        try:
            doc = Document(BytesIO(content))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            if not text_content.strip():
                raise ValueError("No text content found in DOCX document")
            
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Failed to process DOCX: {str(e)}")
            raise
    
    async def _process_email(self, content: bytes) -> str:
        """Extract text from email content"""
        try:
            # Try to decode as UTF-8 first
            email_text = content.decode('utf-8', errors='ignore')
            
            # Parse email
            parser = Parser()
            email_msg = parser.parsestr(email_text)
            
            # Extract basic email info
            text_content = ""
            
            # Add email headers
            headers = ['From', 'To', 'Subject', 'Date']
            for header in headers:
                value = email_msg.get(header)
                if value:
                    text_content += f"{header}: {value}\n"
            
            text_content += "\n--- Email Body ---\n"
            
            # Extract email body
            if email_msg.is_multipart():
                for part in email_msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True)
                        if body:
                            body_text = body.decode('utf-8', errors='ignore')
                            # Parse reply to get main content
                            main_content = EmailReplyParser.parse_reply(body_text)
                            text_content += main_content + "\n"
                    elif part.get_content_type() == "text/html":
                        html_body = part.get_payload(decode=True)
                        if html_body:
                            html_text = html_body.decode('utf-8', errors='ignore')
                            # Convert HTML to text
                            soup = BeautifulSoup(html_text, 'html.parser')
                            body_text = soup.get_text()
                            main_content = EmailReplyParser.parse_reply(body_text)
                            text_content += main_content + "\n"
            else:
                body = email_msg.get_payload(decode=True)
                if body:
                    body_text = body.decode('utf-8', errors='ignore')
                    main_content = EmailReplyParser.parse_reply(body_text)
                    text_content += main_content + "\n"
            
            if not text_content.strip():
                raise ValueError("No text content found in email")
            
            return self._clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Failed to process email: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove special characters that might interfere
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
        
        return text.strip()
    
    def _create_chunks(self, text: str, document_url: str) -> List[DocumentChunk]:
        """Split text into chunks with overlap"""
        chunks = []
        
        # Split into sentences first for better chunking
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        current_chunk = ""
        current_length = 0
        chunk_number = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence would exceed chunk size, create a new chunk
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Create chunk
                chunk_id = self._generate_chunk_id(document_url, chunk_number)
                chunk = DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_id=chunk_id,
                    document_url=document_url,
                    metadata={
                        'chunk_number': chunk_number,
                        'character_count': len(current_chunk),
                        'word_count': len(current_chunk.split()),
                    }
                )
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence
                current_length = len(current_chunk)
                chunk_number += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_length += sentence_length
        
        # Add the last chunk if it has content
        if current_chunk.strip():
            chunk_id = self._generate_chunk_id(document_url, chunk_number)
            chunk = DocumentChunk(
                content=current_chunk.strip(),
                chunk_id=chunk_id,
                document_url=document_url,
                metadata={
                    'chunk_number': chunk_number,
                    'character_count': len(current_chunk),
                    'word_count': len(current_chunk.split()),
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk"""
        words = text.split()
        overlap_words = max(1, self.chunk_overlap // 10)  # Approximate words for overlap
        return " ".join(words[-overlap_words:])
    
    def _generate_chunk_id(self, document_url: str, chunk_number: int) -> str:
        """Generate unique chunk ID"""
        # Create a hash of the document URL for uniqueness
        url_hash = hashlib.md5(document_url.encode()).hexdigest()[:8]
        return f"{url_hash}_{chunk_number}"
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract sections from text based on headers and patterns"""
        sections = {}
        
        # Common section patterns for insurance/legal documents
        section_patterns = [
            r'(?i)(article|section|clause|paragraph|chapter)\s+\d+',
            r'(?i)(coverage|benefits|exclusions|conditions|definitions)',
            r'(?i)(terms and conditions|general provisions|policy conditions)',
            r'(?i)(claims|premium|deductible|co-payment)',
        ]
        
        current_section = "main"
        current_text = ""
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            is_header = False
            for pattern in section_patterns:
                if re.match(pattern, line):
                    # Save previous section
                    if current_text:
                        sections[current_section] = current_text.strip()
                    
                    # Start new section
                    current_section = line
                    current_text = ""
                    is_header = True
                    break
            
            if not is_header:
                current_text += line + "\n"
        
        # Save last section
        if current_text:
            sections[current_section] = current_text.strip()
        
        return sections